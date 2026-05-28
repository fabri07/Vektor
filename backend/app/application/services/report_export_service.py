"""Servicio de exportación de informes de salud financiera.

Genera informes en PDF (reportlab) y DOCX (python-docx) a partir de un
HealthScoreSnapshot + narrativa del AgentHealth.

Usado por: POST /health-scores/{snapshot_id}/export con body {format, narrative}.
"""

from __future__ import annotations

import io
from datetime import datetime
from decimal import Decimal
from typing import Any

from app.persistence.models.score import HealthScoreSnapshot


def _dim_label(key: str) -> str:
    return {
        "cash": "Caja",
        "stock": "Inventario",
        "supplier": "Proveedores",
        "margin": "Márgenes",
        "growth": "Crecimiento",
    }.get(key, key.capitalize())


def _score_label(score: int) -> str:
    if score >= 90:
        return "Excelente"
    if score >= 75:
        return "Bueno"
    if score >= 60:
        return "Regular"
    if score >= 40:
        return "Bajo"
    return "Crítico"


def _format_date(dt: datetime) -> str:
    return dt.strftime("%-d de %B de %Y") if dt else "—"


def _xml_escape(text: str) -> str:
    """Escapa <, >, & — ReportLab interpreta Paragraph como XML mini-markup."""
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _sanitize_narrative(text: str) -> str:
    """Escapa caracteres que rompen ReportLab Paragraph markup y limita longitud."""
    # Limitar a 4000 chars por si llega sin validar (defensa en profundidad)
    return _xml_escape(text[:4000])


def generate_pdf(
    snapshot: HealthScoreSnapshot,
    narrative: str,
    business_name: str,
) -> bytes:
    """Genera un PDF con el informe de salud. Retorna bytes listos para HTTP response."""
    from reportlab.lib import colors  # noqa: PLC0415
    from reportlab.lib.pagesizes import A4  # noqa: PLC0415
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle  # noqa: PLC0415
    from reportlab.lib.units import cm  # noqa: PLC0415
    from reportlab.platypus import (  # noqa: PLC0415
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        rightMargin=2 * cm,
        leftMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "Title",
        parent=styles["Title"],
        fontSize=18,
        spaceAfter=6,
    )
    heading_style = ParagraphStyle(
        "Heading",
        parent=styles["Heading2"],
        fontSize=13,
        spaceBefore=12,
        spaceAfter=4,
    )
    body_style = styles["BodyText"]
    body_style.fontSize = 10
    body_style.leading = 14

    story: list[Any] = []

    # Título (sanitizar business_name por si contiene <, >, &)
    safe_business_name = _xml_escape(business_name)
    story.append(Paragraph(f"Informe de Salud Financiera — {safe_business_name}", title_style))
    story.append(Paragraph(f"Generado el {_format_date(snapshot.snapshot_date)}", body_style))
    story.append(Spacer(1, 0.4 * cm))

    # Score total
    total = int(snapshot.total_score)
    story.append(Paragraph(f"Score general: <b>{total}/100</b> — {_score_label(total)}", heading_style))
    story.append(Spacer(1, 0.3 * cm))

    # Tabla de dimensiones (v2 si score_growth no es None, v1 si es None)
    dims: list[tuple[str, int | None]] = [
        ("Caja", snapshot.score_cash),
        ("Inventario", snapshot.score_stock),
        ("Proveedores", snapshot.score_supplier),
        ("Márgenes", snapshot.score_margin),
    ]
    if snapshot.score_growth is not None:
        dims.append(("Crecimiento", snapshot.score_growth))

    table_data = [["Dimensión", "Score", "Estado"]]
    for label, val in dims:
        if val is not None:
            table_data.append([label, str(val), _score_label(val)])

    table = Table(table_data, colWidths=[5 * cm, 3 * cm, 5 * cm])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1a1a2e")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, 0), 10),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f5f5f5")]),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cccccc")),
        ("ALIGN", (1, 0), (1, -1), "CENTER"),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(table)
    story.append(Spacer(1, 0.5 * cm))

    # Narrativa
    if narrative:
        safe_narrative = _sanitize_narrative(narrative)
        story.append(Paragraph("Análisis ejecutivo", heading_style))
        for para in safe_narrative.split("\n\n"):
            if para.strip():
                story.append(Paragraph(para.strip(), body_style))
                story.append(Spacer(1, 0.2 * cm))

    doc.build(story)
    return buf.getvalue()


def generate_docx(
    snapshot: HealthScoreSnapshot,
    narrative: str,
    business_name: str,
) -> bytes:
    """Genera un DOCX con el informe de salud. Retorna bytes listos para HTTP response."""
    from docx import Document  # noqa: PLC0415
    from docx.shared import Pt, RGBColor  # noqa: PLC0415

    doc = Document()

    # Título
    title = doc.add_heading(f"Informe de Salud Financiera — {business_name}", level=1)
    title.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    doc.add_paragraph(f"Generado el {_format_date(snapshot.snapshot_date)}")

    total = int(snapshot.total_score)
    doc.add_heading(f"Score general: {total}/100 — {_score_label(total)}", level=2)

    # Tabla dimensiones
    dims: list[tuple[str, int | None]] = [
        ("Caja", snapshot.score_cash),
        ("Inventario", snapshot.score_stock),
        ("Proveedores", snapshot.score_supplier),
        ("Márgenes", snapshot.score_margin),
    ]
    if snapshot.score_growth is not None:
        dims.append(("Crecimiento", snapshot.score_growth))

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Dimensión"
    hdr[1].text = "Score"
    hdr[2].text = "Estado"
    for label, val in dims:
        if val is not None:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = str(val)
            row[2].text = _score_label(val)

    doc.add_paragraph()

    # Narrativa (python-docx no interpreta XML, pero limitamos longitud igual)
    if narrative:
        safe_narrative = narrative[:4000]
        doc.add_heading("Análisis ejecutivo", level=2)
        for para in safe_narrative.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
